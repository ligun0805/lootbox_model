import os
import random
import datetime
import argparse
import openpyxl
import pandas as pd
from docx import Document

# Парсер аргументов: число открытий
parser = argparse.ArgumentParser(
    description="Генерация статистики выпадений лутбоксов"
)
parser.add_argument(
    '-n', '--opens', type=int, default=100_000,
    help="Сколько открытий смоделировать (по умолчанию 100000)"
)
args = parser.parse_args()
N = args.opens

# Генерация папки под выводы с таймстампом
BASE_OUT = 'outputs'
ts = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
out_dir = os.path.join(BASE_OUT, ts)
os.makedirs(out_dir, exist_ok=True)
xlsx_path = os.path.join(out_dir, 'lootbox_model.xlsx')
docx_path = os.path.join(out_dir, 'lootbox_model.docx')

# 1) Базовые вероятности
base_probs = {
    'Common':    60.0,
    'Rare':      25.0,
    'Epic':      10.0,
    'Legendary':  5.0,
}
rates_df = pd.DataFrame(
    [(k, f"{v:.1f}%") for k, v in base_probs.items()],
    columns=['Rarity', 'Probability']
)

# 2) Pity System для Rare/Epic/Legendary
pity_thresholds = {
    'Rare': 10,
    'Epic': 50,
    'Legendary': 200,
}
max_pity = max(pity_thresholds.values())
pity_records = []
for i in range(1, max_pity + 1):
    rec = {'Open': i}
    for cat, base in base_probs.items():
        if cat in pity_thresholds:
            th = pity_thresholds[cat]
            inc = (100.0 - base) / th
            chance = min(base + inc * (i - 1), 100.0)
        else:
            chance = base
        rec[f'{cat}ChancePercent'] = round(chance, 2)
    pity_records.append(rec)
pity_df = pd.DataFrame(pity_records)

# 3) Симуляция выпадений
categories = list(base_probs.keys())
weights = [base_probs[c] for c in categories]
results = random.choices(categories, weights=weights, k=N)
counts = {c: results.count(c) for c in categories}
sim_df = pd.DataFrame(
    [(c, counts[c], round(counts[c] / N * 100, 2)) for c in categories],
    columns=['Category', 'Count', 'Percent']
)

# 4) Экономическая модель
scenarios = [
    ('Optimistic',  10000, 20),
    ('Average',      5000, 15),
    ('Pessimistic',  2000, 10),
]
econ = []
for name, players, opens in scenarios:
    revenue = players * opens * 2
    profit = revenue * 0.5
    arpu = revenue / players
    ltv = arpu * 3
    econ.append((name, players, opens,
                 revenue, profit,
                 round(arpu, 2), round(ltv, 2)))
econ_df = pd.DataFrame(
    econ,
    columns=['Scenario', 'Players', 'OpensPerPlayer',
             'Revenue', 'Profit', 'ARPU', 'LTV']
)

# 5) Сохранение в Excel с графиками
with pd.ExcelWriter(xlsx_path, engine='xlsxwriter') as writer:
    rates_df.to_excel(writer, sheet_name='DropRates', index=False)
    pity_df.to_excel(writer, sheet_name='PitySystem', index=False)
    sim_df.to_excel(writer, sheet_name='SimulationResults', index=False)
    econ_df.to_excel(writer, sheet_name='EconomicModel', index=False)

    workbook  = writer.book
    ws_pity   = writer.sheets['PitySystem']
    ws_sim    = writer.sheets['SimulationResults']

    # График для PitySystem (Common, Rare, Epic, Legendary)
    chart_pity = workbook.add_chart({'type': 'line'})
    max_row = len(pity_df)
    # серии: столбцы 1…4 в порядке ['Common','Rare','Epic','Legendary']
    for idx, cat in enumerate(['Common','Rare','Epic','Legendary'], start=1):
        chart_pity.add_series({
            'name':       cat,
            'categories': ['PitySystem', 1, 0, max_row, 0],  # Open
            'values':     ['PitySystem', 1, idx, max_row, idx],
        })
    chart_pity.set_title({'name': 'Pity System Dynamics'})
    chart_pity.set_x_axis({'name': 'Open Number'})
    chart_pity.set_y_axis({'name': 'Chance Percent'})
    ws_pity.insert_chart('G2', chart_pity, {'x_scale': 1.2, 'y_scale': 1.5})
    # График для SimulationResults
    chart_sim = workbook.add_chart({'type': 'column'})
    chart_sim.add_series({
        'name':       'Simulation %',
        'categories': ['SimulationResults', 1, 0, len(sim_df), 0],
        'values':     ['SimulationResults', 1, 2, len(sim_df), 2],
    })
    chart_sim.set_title({'name': 'Simulation Results (%)'})
    chart_sim.set_x_axis({'name': 'Category'})
    chart_sim.set_y_axis({'name': 'Percent'})
    ws_sim.insert_chart('E2', chart_sim, {'x_scale': 1.2, 'y_scale': 1.5})

# 6) Сборка Word‑документа

doc = Document()
doc.add_heading('Модель выпадения из лутбоксов', level=1)

doc.add_heading('1. Базовые шансы выпадения', level=2)
tbl = doc.add_table(rows=1, cols=2)
hdr = tbl.rows[0].cells
hdr[0].text, hdr[1].text = 'Категория', 'Вероятность'
for k, v in base_probs.items():
    row = tbl.add_row().cells
    row[0].text = k
    row[1].text = f"{v:.1f}%"

doc.add_heading('2. Pity System', level=2)
doc.add_paragraph(
    'Гарант через Rare=10, Epic=50, Legendary=200 открытий. '
    'Шанс растёт линейно до 100%.'
)

doc.add_heading('3. Результаты симуляции', level=2)
doc.add_paragraph(
    f'Симуляция {N} открытий. В разделе SimulationResults '  
    'Excel-файла представлены фактические частоты.'
)

doc.add_heading('4. Экономическая модель', level=2)
tbl2 = doc.add_table(rows=1, cols=7)
hdr2 = tbl2.rows[0].cells
for i, title in enumerate(['Scenario', 'Players', 'Opens', 'Revenue', 'Profit', 'ARPU', 'LTV']):
    hdr2[i].text = title
for row in econ:
    cells = tbl2.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = str(val)

doc.save(docx_path)

print(f"✅ Симуляция {N} запусков завершена.")
print(f"✅ Файлы со схемами и графиками сохранены в: {out_dir}")
